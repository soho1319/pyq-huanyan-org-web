import os
from docx import Document
from pathlib import Path

src_dir = Path(r"D:\ObsidianVaults\内容营销朋友圈文案\doc")
out_dir = Path(r"D:\ObsidianVaults\内容营销朋友圈文案\doc")

def docx_to_md(docx_path):
    doc = Document(docx_path)
    md_lines = []

    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style else ""
        text = para.text.strip()

        if not text:
            md_lines.append("")
            continue

        # 标题处理
        if style_name.startswith("heading") or "标题" in style_name:
            level = 1
            if "2" in style_name:
                level = 2
            elif "3" in style_name:
                level = 3
            md_lines.append(f"{'#' * level} {text}")
        # 列表处理
        elif style_name.startswith("list") or "列表" in style_name:
            md_lines.append(f"- {text}")
        # 普通段落
        else:
            md_lines.append(text)

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("")

    return "\n".join(md_lines)

# 转换所有docx文件
count = 0
for docx_file in sorted(src_dir.glob("*.docx")):
    md_content = docx_to_md(docx_file)
    md_name = docx_file.stem + ".md"
    md_path = out_dir / md_name

    if md_path.exists():
        md_path.unlink()

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    count += 1
    print(f"转换: {docx_file.name} -> {md_name}")

print(f"\n完成! 共转换 {count} 个文件")